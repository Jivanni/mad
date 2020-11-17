from docx import Document
import re, pandas, copy, json

def cleanup(string):
    return re.sub(r"\"","",string).strip()


def display_doc(doc):
    document = Document(doc)
    for par in document.paragraphs:
        print(par.text)


def fill_mad_doc(doc, user_data):
    d = copy.deepcopy(doc)
    for par in d.paragraphs:
        par.text = par.text.format(user = user_data)
    try: d.save("docs/{user[nome_istituto]}.docx".format(user = user_data))
    except: d.save("docs/domanda messa a disposizione.docx")


def flitered_list(lista,**kwargs):
    email = lista.loc[:, ["DENOMINAZIONESCUOLA",'INDIRIZZOEMAILSCUOLA',"INDIRIZZOPECSCUOLA","REGIONE","PROVINCIA"]]
    dic = email.to_dict(orient="index")
    filtered_dic = {key:value for key, value in dic.items() if dic[key]["REGIONE"] == kwargs["REGIONE"] and dic[key]["PROVINCIA"] == kwargs["PROVINCIA"]}
    finale = []
    for scuola_id in filtered_dic:
        email = ""
        if filtered_dic[scuola_id]['INDIRIZZOPECSCUOLA'] != "Non Disponibile":
            email = filtered_dic[scuola_id]['INDIRIZZOPECSCUOLA']
        else: email = filtered_dic[scuola_id]['INDIRIZZOEMAILSCUOLA']
        finale.append([cleanup(filtered_dic[scuola_id]["DENOMINAZIONESCUOLA"]),email])
    return finale

def create_docs(lista, document, user_data, scuole_target):
    for i in range(len(scuole_target)):
        user_data["nome_istituto"] = scuole_target[i][0]
        fill_mad_doc(document, user_data)

